"""S147-3 toolbar bundle unit coverage for ``OfficeDocExportService`` —
Print/PDF/DOCX/Markdown export orchestration (access resolution, format
dispatch, and the shared asset-bytes resolver PDF/DOCX embed images
through). Against in-memory fakes with an injected ``_FakePdfService``
standing in for the CORE ``pdf_service`` seam — mirrors
``test_sheet_editor_service.py``'s established pattern for this plugin. No
DB, no Flask app.
"""
import io
import json
import uuid

import docx
import pytest

from plugins.office.office.models.office_document import DOC_TYPE_FILE, DOC_TYPE_TEXT
from plugins.office.office.models.office_node import NODE_KIND_DOCUMENT
from plugins.office.office.services.access_resolver import ACCESS_OWNER
from plugins.office.office.services.doc_export import (
    DOCX_MIME_TYPE,
    MARKDOWN_MIME_TYPE,
    PDF_MIME_TYPE,
    OfficeDocExportService,
)
from plugins.office.office.services.exceptions import OfficeDocExportFormatError

#: A real (if trivial) 1x1 PNG — python-docx fully parses the PNG chunk
#: structure to size the embedded picture, so a bare magic-number prefix
#: (sufficient for ``MimeSniffer``) is not enough here.
PNG_BYTES = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\xcf\xc0"
    b"\x00\x00\x03\x01\x01\x00\x18\xdd\x8d\xb0\x00\x00\x00\x00IEND\xaeB`\x82"
)


class _Node:
    def __init__(self, node_id, owner_user_id, name="My Doc", kind=NODE_KIND_DOCUMENT):
        self.id = node_id
        self.owner_user_id = owner_user_id
        self.name = name
        self.kind = kind


class _Document:
    def __init__(
        self, document_id, doc_type=DOC_TYPE_TEXT, mime_type="application/json"
    ):
        self.id = document_id
        self.doc_type = doc_type
        self.mime_type = mime_type


class FakeNodeRepository:
    def __init__(self, nodes_by_id):
        self._nodes = nodes_by_id

    def find_by_id(self, node_id):
        return self._nodes.get(node_id)

    def find_by_id_for_owner(self, node_id, owner_user_id):
        # Mirrors the real repository accepting either a UUID or its string
        # form (SQLAlchemy coerces a str into the UUID column type) — the
        # asset resolver strips the ``office-asset:`` prefix into a string.
        node = self._nodes.get(node_id) or self._nodes.get(uuid.UUID(str(node_id)))
        if node is None or node.owner_user_id != owner_user_id:
            return None
        return node


class FakeDocumentRepository:
    def __init__(self, documents_by_node_id):
        self._documents = documents_by_node_id

    def find_by_node_id(self, node_id):
        return self._documents.get(node_id)


class FakeAccessResolver:
    def __init__(self, access_by_user):
        self._access_by_user = access_by_user

    def resolve(self, node_id, *, user_id=None, share_token=None):
        return self._access_by_user.get(user_id)


class FakeDocumentService:
    def __init__(self, content_by_document_id):
        self._content_by_document_id = content_by_document_id

    def get_content_for_node(self, node, document, version_no=None):
        return node, document, None, self._content_by_document_id[document.id]


class _FakePdfService:
    """Stands in for the CORE ``pdf_service`` — records what the export
    service asked it to render (mirrors ``test_sheet_editor_service.py``)."""

    def __init__(self) -> None:
        self.registered_paths = []
        self.rendered = []

    def register_plugin_template_path(self, path) -> None:
        self.registered_paths.append(path)

    def render(self, template_name, context) -> bytes:
        self.rendered.append((template_name, context))
        return b"%PDF-1.4 fake"


def _wiring(owner_id, content_model, *, asset_content=None):
    doc_node_id = uuid.uuid4()
    doc_node = _Node(doc_node_id, owner_id, name="Quarterly Report")
    doc_document = _Document(uuid.uuid4())
    nodes = {doc_node_id: doc_node}
    documents = {doc_node_id: doc_document}
    content_by_document_id = {
        doc_document.id: json.dumps(content_model).encode("utf-8")
    }

    if asset_content is not None:
        asset_node_id = uuid.uuid4()
        nodes[asset_node_id] = _Node(asset_node_id, owner_id, name="cat.png")
        asset_document = _Document(
            uuid.uuid4(), doc_type=DOC_TYPE_FILE, mime_type="image/png"
        )
        documents[asset_node_id] = asset_document
        content_by_document_id[asset_document.id] = asset_content

    pdf_service = _FakePdfService()
    service = OfficeDocExportService(
        node_repository=FakeNodeRepository(nodes),
        document_repository=FakeDocumentRepository(documents),
        document_service=FakeDocumentService(content_by_document_id),
        access_resolver=FakeAccessResolver({owner_id: ACCESS_OWNER}),
        pdf_service=pdf_service,
    )
    asset_node_id = asset_node_id if asset_content is not None else None
    return service, doc_node_id, pdf_service, asset_node_id


def _paragraph(text: str) -> dict:
    return {"type": "paragraph", "content": [{"type": "text", "text": text}]}


def test_markdown_export_returns_the_rendered_bytes_mimetype_and_filename():
    owner_id = uuid.uuid4()
    model = {"type": "doc", "content": [_paragraph("hello world")]}
    service, node_id, _pdf, _asset = _wiring(owner_id, model)

    data, mimetype, filename = service.export_document(owner_id, node_id, "md")

    assert data == b"hello world\n"
    assert mimetype == MARKDOWN_MIME_TYPE
    assert filename == "Quarterly Report.md"


def test_docx_export_returns_valid_docx_bytes_with_the_expected_title():
    owner_id = uuid.uuid4()
    model = {"type": "doc", "content": [_paragraph("hello world")]}
    service, node_id, _pdf, _asset = _wiring(owner_id, model)

    data, mimetype, filename = service.export_document(owner_id, node_id, "docx")

    assert mimetype == DOCX_MIME_TYPE
    assert filename == "Quarterly Report.docx"
    document = docx.Document(io.BytesIO(data))
    assert document.core_properties.title == "Quarterly Report"
    assert "hello world" in [p.text for p in document.paragraphs]


def test_docx_export_embeds_an_image_via_the_shared_asset_resolver():
    owner_id = uuid.uuid4()
    # Build with a placeholder src, then re-point it at the asset id
    # ``_wiring`` mints — the model has to reference an id that already
    # exists in the same service's repositories.
    model = {
        "type": "doc",
        "content": [{"type": "image", "attrs": {"src": "office-asset:PLACEHOLDER"}}],
    }
    service, node_id, _pdf, asset_node_id = _wiring(
        owner_id, model, asset_content=PNG_BYTES
    )
    model["content"][0]["attrs"]["src"] = f"office-asset:{asset_node_id}"
    doc_document = service._document_repository.find_by_node_id(node_id)
    service._document_service._content_by_document_id[doc_document.id] = json.dumps(
        model
    ).encode("utf-8")

    data, _mimetype, _filename = service.export_document(owner_id, node_id, "docx")

    document = docx.Document(io.BytesIO(data))
    assert len(document.inline_shapes) == 1


def test_pdf_export_renders_via_the_core_pdf_service_and_returns_its_bytes():
    owner_id = uuid.uuid4()
    model = {"type": "doc", "content": [_paragraph("hello world")]}
    service, node_id, pdf_service, _asset = _wiring(owner_id, model)

    data, mimetype, filename = service.export_document(owner_id, node_id, "pdf")

    assert data == b"%PDF-1.4 fake"
    assert mimetype == PDF_MIME_TYPE
    assert filename == "Quarterly Report.pdf"
    assert len(pdf_service.rendered) == 1
    template_name, context = pdf_service.rendered[0]
    assert template_name == "office_doc.html"
    assert context["title"] == "Quarterly Report"
    assert "hello world" in context["body_html"]
    assert len(pdf_service.registered_paths) == 1  # self-heals the template path


def test_export_rejects_an_unsupported_format():
    owner_id = uuid.uuid4()
    model = {"type": "doc", "content": [_paragraph("hello")]}
    service, node_id, _pdf, _asset = _wiring(owner_id, model)

    with pytest.raises(OfficeDocExportFormatError):
        service.export_document(owner_id, node_id, "csv")
