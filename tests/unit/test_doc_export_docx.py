"""S147-3 toolbar bundle unit coverage for ``render_docx`` — round-trips the
produced ``.docx`` bytes back through ``python-docx`` to assert on the real
document tree (no DB, no Flask app).
"""
import io

import docx
import pytest

from plugins.office.office.services.doc_export_docx import render_docx


def _paragraph(text: str, marks=None) -> dict:
    return {
        "type": "paragraph",
        "content": [{"type": "text", "text": text, "marks": marks or []}],
    }


def _open(data: bytes) -> docx.Document:
    return docx.Document(io.BytesIO(data))


def test_sets_the_document_title_from_the_node_name():
    model = {"type": "doc", "content": [_paragraph("hello")]}
    data = render_docx(model, "My Document")
    assert _open(data).core_properties.title == "My Document"


def test_renders_a_heading_and_a_paragraph_as_separate_blocks():
    model = {
        "type": "doc",
        "content": [
            {
                "type": "heading",
                "attrs": {"level": 1},
                "content": [{"type": "text", "text": "Title"}],
            },
            _paragraph("body text"),
        ],
    }
    document = _open(render_docx(model, "doc"))
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert "Title" in texts
    assert "body text" in texts


def test_bold_italic_underline_and_strike_marks_apply_to_the_run():
    model = {
        "type": "doc",
        "content": [
            _paragraph("b", [{"type": "bold"}]),
            _paragraph("i", [{"type": "italic"}]),
            _paragraph("u", [{"type": "underline"}]),
            _paragraph("s", [{"type": "strike"}]),
        ],
    }
    document = _open(render_docx(model, "doc"))
    runs_by_text = {p.runs[0].text: p.runs[0] for p in document.paragraphs if p.runs}
    assert runs_by_text["b"].bold is True
    assert runs_by_text["i"].italic is True
    assert runs_by_text["u"].underline is True
    assert runs_by_text["s"].font.strike is True


def test_a_table_round_trips_its_cell_text():
    model = {
        "type": "doc",
        "content": [
            {
                "type": "table",
                "content": [
                    {
                        "type": "tableRow",
                        "content": [
                            {"type": "tableCell", "content": [_paragraph("A1")]},
                            {"type": "tableCell", "content": [_paragraph("B1")]},
                        ],
                    }
                ],
            }
        ],
    }
    document = _open(render_docx(model, "doc"))
    assert len(document.tables) == 1
    table = document.tables[0]
    assert table.cell(0, 0).text == "A1"
    assert table.cell(0, 1).text == "B1"


def test_an_image_is_embedded_via_the_asset_bytes_resolver():
    tiny_png = (
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
        b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\xcf\xc0"
        b"\x00\x00\x03\x01\x01\x00\x18\xdd\x8d\xb0\x00\x00\x00\x00IEND\xaeB`\x82"
    )
    model = {
        "type": "doc",
        "content": [{"type": "image", "attrs": {"src": "office-asset:asset-1"}}],
    }
    resolver = lambda asset_id: (tiny_png, "image/png")  # noqa: E731
    document = _open(render_docx(model, "doc", resolver))
    assert len(document.inline_shapes) == 1


def test_an_image_the_resolver_cannot_resolve_is_silently_omitted():
    model = {
        "type": "doc",
        "content": [{"type": "image", "attrs": {"src": "office-asset:missing"}}],
    }
    document = _open(render_docx(model, "doc", lambda asset_id: None))
    assert len(document.inline_shapes) == 0


def test_a_textstyle_font_family_sets_the_run_font_name():
    model = {
        "type": "doc",
        "content": [
            _paragraph(
                "styled", [{"type": "textStyle", "attrs": {"fontFamily": "Georgia"}}]
            )
        ],
    }
    document = _open(render_docx(model, "doc"))
    run = next(p.runs[0] for p in document.paragraphs if p.runs)
    assert run.font.name == "Georgia"


def test_a_textstyle_font_family_with_a_quoted_multi_word_stack_uses_the_first_name():
    model = {
        "type": "doc",
        "content": [
            _paragraph(
                "styled",
                [
                    {
                        "type": "textStyle",
                        "attrs": {"fontFamily": "'Times New Roman', serif"},
                    }
                ],
            )
        ],
    }
    document = _open(render_docx(model, "doc"))
    run = next(p.runs[0] for p in document.paragraphs if p.runs)
    assert run.font.name == "Times New Roman"


@pytest.mark.parametrize(
    "font_size,expected_points",
    [("16px", pytest.approx(12.0, abs=0.1)), ("14pt", 14.0)],
)
def test_a_textstyle_font_size_converts_to_points(font_size, expected_points):
    model = {
        "type": "doc",
        "content": [
            _paragraph(
                "styled", [{"type": "textStyle", "attrs": {"fontSize": font_size}}]
            )
        ],
    }
    document = _open(render_docx(model, "doc"))
    run = next(p.runs[0] for p in document.paragraphs if p.runs)
    assert run.font.size.pt == expected_points
