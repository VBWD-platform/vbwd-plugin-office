"""Content-model -> ``.docx`` rendering for VBWD Docs export (S147-3
toolbar bundle), via ``python-docx`` (``plugins/office/requirements.txt``).

Same "good-enough round-tripping, not Word parity" honesty the epic already
applies to import/export elsewhere: a table cell keeps its plain text (not
per-run marks), a blockquote becomes an indented paragraph, and a horizontal
rule becomes a text divider rather than a real border — all documented
trade-offs, never a silent loss of the whole export (S147-3's own escape
hatch: ship less, but say so).
"""
from __future__ import annotations

import io
import re
from typing import Any, Callable, Dict, List, Optional

#: Asset bytes resolver: ``asset_node_id -> (bytes, mime_type)`` or ``None``.
AssetBytesResolver = Optional[Callable[[str], Optional[tuple]]]

_HORIZONTAL_RULE_TEXT = "—" * 24
_BLOCKQUOTE_INDENT_INCHES = 0.5
_DEFAULT_IMAGE_WIDTH_INCHES = 4.0


def render_docx(
    content_model: Dict[str, Any],
    title: str,
    asset_bytes_resolver: AssetBytesResolver = None,
) -> bytes:
    import docx

    document = docx.Document()
    document.core_properties.title = title
    _render_blocks(document, content_model.get("content") or [], asset_bytes_resolver)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _render_blocks(
    document_or_cell, nodes: List[Dict[str, Any]], resolver: AssetBytesResolver
) -> None:
    for node in nodes:
        _render_block(document_or_cell, node, resolver)


def _render_block(
    document_or_cell, node: Dict[str, Any], resolver: AssetBytesResolver
) -> None:
    node_type = node.get("type")
    children = node.get("content") or []
    if node_type == "paragraph":
        _render_inline(document_or_cell.add_paragraph(), children)
    elif node_type == "heading":
        level = min(max(int((node.get("attrs") or {}).get("level", 1)), 1), 6)
        _render_inline(document_or_cell.add_heading(level=level), children)
    elif node_type == "bulletList":
        _render_list(document_or_cell, children, "List Bullet")
    elif node_type == "orderedList":
        _render_list(document_or_cell, children, "List Number")
    elif node_type == "blockquote":
        _render_blockquote(document_or_cell, children, resolver)
    elif node_type == "codeBlock":
        _render_code_block(document_or_cell, node)
    elif node_type == "table":
        _render_table(document_or_cell, node, resolver)
    elif node_type == "image":
        _render_image(document_or_cell, node, resolver)
    elif node_type == "horizontalRule":
        document_or_cell.add_paragraph(_HORIZONTAL_RULE_TEXT)
    else:
        _render_blocks(document_or_cell, children, resolver)  # the "doc" root


def _render_list(document, items: List[Dict[str, Any]], style_name: str) -> None:
    for item in items:
        for child in item.get("content") or []:
            if child.get("type") == "paragraph":
                _render_inline(
                    document.add_paragraph(style=style_name), child.get("content") or []
                )
            else:
                _render_block(document, child, None)


def _render_blockquote(
    document, children: List[Dict[str, Any]], resolver: AssetBytesResolver
) -> None:
    from docx.shared import Inches

    for child in children:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(_BLOCKQUOTE_INDENT_INCHES)
        run = paragraph.add_run(_extract_plain_text_recursive(child))
        run.italic = True


def _render_code_block(document, node: Dict[str, Any]) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(_extract_plain_text(node))
    run.font.name = "Consolas"


def _render_table(document, node: Dict[str, Any], resolver: AssetBytesResolver) -> None:
    rows = node.get("content") or []
    if not rows:
        return
    column_count = max((len(row.get("content") or []) for row in rows), default=0)
    if column_count == 0:
        return
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index, cell in enumerate(row.get("content") or []):
            table.cell(row_index, column_index).text = _extract_plain_text_recursive(
                cell
            )


def _render_image(document, node: Dict[str, Any], resolver: AssetBytesResolver) -> None:
    from plugins.office.office.services.doc_content import IMAGE_SRC_PREFIX

    if resolver is None:
        return
    src = (node.get("attrs") or {}).get("src", "")
    if not isinstance(src, str) or not src.startswith(IMAGE_SRC_PREFIX):
        return
    resolved = resolver(src[len(IMAGE_SRC_PREFIX) :])
    if resolved is None:
        return
    image_bytes, _mime_type = resolved

    from docx.shared import Inches

    document.add_picture(
        io.BytesIO(image_bytes), width=Inches(_DEFAULT_IMAGE_WIDTH_INCHES)
    )


def _render_inline(paragraph, nodes: List[Dict[str, Any]]) -> None:
    for node in nodes:
        if node.get("type") != "text":
            continue
        run = paragraph.add_run(node.get("text", ""))
        for mark in node.get("marks") or []:
            _apply_mark(run, mark)


def _apply_mark(run, mark: Dict[str, Any]) -> None:
    mark_type = mark.get("type")
    if mark_type == "bold":
        run.bold = True
    elif mark_type == "italic":
        run.italic = True
    elif mark_type == "underline":
        run.underline = True
    elif mark_type == "strike":
        run.font.strike = True
    elif mark_type == "code":
        run.font.name = "Consolas"
    elif mark_type == "textStyle":
        _apply_text_style(run, mark.get("attrs") or {})


def _apply_text_style(run, attrs: Dict[str, Any]) -> None:
    font_family = attrs.get("fontFamily")
    if font_family:
        # A Tiptap font stack may be "'Times New Roman', serif" — a .docx
        # run takes exactly one family name, so the first (most specific,
        # quote-stripped) entry in the stack is what gets applied.
        run.font.name = font_family.split(",")[0].strip().strip("'\"")
    font_size = attrs.get("fontSize")
    if font_size:
        points = _font_size_to_points(font_size)
        if points is not None:
            from docx.shared import Pt

            run.font.size = Pt(points)


#: 1rem/1em is treated as the browser default body font size for a
#: good-enough (not pixel-perfect) points conversion — the same "honest
#: trade-off" this module's docstring already applies elsewhere.
_ASSUMED_ROOT_FONT_SIZE_PX = 16.0
_CSS_PIXELS_PER_POINT = 96.0 / 72.0


def _font_size_to_points(font_size: str) -> Optional[float]:
    match = re.match(r"^(\d+(?:\.\d+)?)(px|pt|em|rem)$", font_size)
    if not match:
        return None
    magnitude, unit = float(match.group(1)), match.group(2)
    if unit == "pt":
        return magnitude
    if unit == "px":
        return magnitude / _CSS_PIXELS_PER_POINT
    return (magnitude * _ASSUMED_ROOT_FONT_SIZE_PX) / _CSS_PIXELS_PER_POINT


def _extract_plain_text(node: Dict[str, Any]) -> str:
    parts: List[str] = []
    for child in node.get("content") or []:
        if child.get("type") == "text":
            parts.append(child.get("text", ""))
        elif child.get("type") == "hardBreak":
            parts.append("\n")
    return "".join(parts)


def _extract_plain_text_recursive(node: Dict[str, Any]) -> str:
    """Like :func:`_extract_plain_text` but descends through block wrappers
    (paragraph/listItem/…) — used for table cells and blockquote lines,
    where python-docx keeps ONE paragraph's worth of plain text per cell/
    line rather than the full nested block structure (documented fidelity
    trade-off, see the module docstring)."""
    if node.get("type") == "text":
        return node.get("text", "")
    parts = [
        _extract_plain_text_recursive(child) for child in node.get("content") or []
    ]
    return "".join(parts)
